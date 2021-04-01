from docx import Document

Dictionary = {"pdf": "ocean"}

import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def main(board,board_sol):
    filename = "template.docx"
    doc = Document(filename)
    rows = ['a','b','c','d','e','f','g','h','i']
    cols = ['1','2','3','4','5','6','7','8','9']
    # question
    for i in range(9):
        for j in range(9):
            regex1 = re.compile(rows[i]+cols[j])
            replace1 = str(board[i][j])
            if board[i][j] == 0:
                replace1 = " "
            docx_replace_regex(doc, regex1 , replace1)
    # solution
    sol_rows = ['j','k','l','m','n','o','p','q','r']
    for i in range(9):
        for j in range(9):
            regex2 = re.compile(sol_rows[i]+cols[j])
            replace2 = str(board_sol[i][j])
            if board_sol[i][j] == 0:
                replace2 = " "
            docx_replace_regex(doc, regex2 , replace2)
    doc.save('result1.docx')
